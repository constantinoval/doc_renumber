# %%
from colorama import Fore, Back, init
import docx
import re
from doctools_lib import replace_text_in_runs, unpack_ref, pack_ref, paragraph_iterator


def get_fig_params(input_str):
    p1 = re.compile(r'<f>(.*)<\\f>')
    start = re.compile(r'start\s*=\s*([\d]+)')
    prefix = re.compile(r'prefix\s*=\s*([\d.]+)')
    m1 = p1.search(input_str)
    if not m1:
        return None
    kws = {}
    kws['prefix'] = prefix.search(m1[1])[1] if prefix.search(m1[1]) else ''
    kws['start'] = int(start.search(m1[1])[1]) if start.search(m1[1]) else 1
    return kws


fig_in_text = re.compile(
    r'(рис((.)|(ун((ке)|(ок)|(ка)|(ках))))\s+)(([\s,и]*[\d.\-–]+)*[^\D])', re.IGNORECASE)
stop_tag = re.compile(r'<f?stop>')
continue_tag = re.compile(r'<f?continue>')


def analize_figures(document, prefix='', start=1):
    pat = re.compile(
        r'^\s*(Рис\.|Рисунок)\s+(?P<num>[\d.]+?)\s*([–-].*|[. ]*$)')
    rez = {}
    do_analysis = True
    for p in paragraph_iterator(document):
        if stop_tag.search(p.text):
            do_analysis = False
        if continue_tag.search(p.text):
            do_analysis = True
        if not do_analysis:
            continue
        kw = get_fig_params(p.text)
        if kw:
            start = kw['start']
            prefix = kw['prefix']
            continue
        m = pat.match(p.text)
        if m:
            rez[m.group('num')] = [prefix, start]
            start += 1
    return rez


def renumber_figures(inp, output, prefix, start):
    init(autoreset=True)
    figs = analize_figures(docx.Document(inp), prefix, start)
    doc = docx.Document(inp)
    do_replacement = True
    for p in paragraph_iterator(doc):  # doc.paragraphs:
        if stop_tag.search(p.text):
            do_replacement = False
        if continue_tag.search(p.text):
            do_replacement = True
        if not do_replacement:
            continue
        ss = fig_in_text.search(p.text)
        while ss:
            # body = ss[1]
            _, _, nn = unpack_ref(ss[10])
            for n in nn:
                if not n in figs:
                    print(Fore.RED+f'Figure {n} is not found in document')
                    nn.remove(n)
            if nn:
                new_str = pack_ref([figs[n][1] for n in nn], figs[nn[0]][0])
                print(ss[0], '->', new_str)
                replace_text_in_runs(p.runs, ss.span(10)[0],
                                     ss.span(10)[1], new_str)
            ss = fig_in_text.search(p.text, ss.span(10)[1]+1)
    doc.save(output)


if __name__ == '__main__':
    from argparse import ArgumentParser
    parser = ArgumentParser()
    parser.add_argument('input', help='Input docx file')
    parser.add_argument('--start', '-n', default=1,
                        help='Start number for figures numbering', type=int)
    parser.add_argument('--prefix', '-p', default='',
                        help='Default prefix for numbers')
    parser.add_argument(
        '--output', '-o', default='rez.docx', help='Output file')
    args = parser.parse_args()
    renumber_figures(args.input, args.output, args.prefix, args.start)
