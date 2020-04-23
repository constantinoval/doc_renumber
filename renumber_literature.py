from difflib import SequenceMatcher
from itertools import chain
import docx
import re
import pandas as pd
from os.path import exists
from doctools_lib import pack_ref, unpack_ref, replace_text_in_runs, paragraph_iterator
import numpy as np


def fix_dublicates(data):
    dubs = {}  # OrderedDict()
    N = len(data)
    data.reset_index(inplace=True)
    for i in range(N-1):
        for j in range(i+1, N):
            if j == i:
                continue
            r = SequenceMatcher(lambda x: x in ' \t\n',
                                data.iloc[i]['title'],
                                data.iloc[j]['title']).ratio()
            if r > 0.95:
                ii = data.new_n[i]
                jj = data.new_n[j]
                if not ii in dubs:
                    dubs[ii] = []
                dubs[ii].append(jj)
    if not dubs:
        data.set_index(keys='n', inplace=True)
        return data
    k = list(dubs.keys())
    k.reverse()
    for kk in k:
        data.loc[data.new_n.map(lambda x: x in dubs[kk]),
                 'new_n'] = kk
    data.sort_values(by='new_n', inplace=True)
    data['d'] = data.new_n.diff()-1.
    data['d'] = data['d'].map(lambda x: max(0, x))
    data.fillna(0, inplace=True)
    data['d'] = data['d'].cumsum().astype(int)
    data['new_n'] -= data['d']
    data.drop(columns='d', inplace=True)
    data.set_index(keys='n', inplace=True)
    return data


t = re.compile(r'\[([0123456789\-,;\s]*)\]')


def renumber_refs(inp, output, refs_input=None, refs_output='new_refs.xlsx', start=1):
    if refs_input == None:
        if exists(inp[:-4]+'xlsx'):
            refs_input = inp[:-4]+'xlsx'
    if refs_input:
        lit = pd.read_excel(refs_input)
        lit.n = lit.n.astype(dtype=np.dtype(int))
        lit.set_index(keys='n', inplace=True)
    doc = docx.Document(inp)
    new_n = {}
    for p in paragraph_iterator(doc):  # doc.paragraphs:
        for ss in t.findall(p.text):
            for sss in unpack_ref(ss)[1]:
                if refs_input and (not (int(sss) in lit.index)):
                    new_n[sss] = -1
                    lit.loc[int(sss)] = 'Not found'
                    continue
                if not sss in new_n:
                    new_n[sss] = start
                    start += 1

    r = list(map(int, new_n.keys()))
    if refs_input:
        # for rr in r:
        #     if not rr in lit:
        #         lit.loc[rr] = str(rr)
        lit = lit.loc[r]
        lit['new_n'] = list(new_n.values())
        lit = fix_dublicates(lit)
    else:
        lit = pd.DataFrame(
            data={'n': list(new_n.keys()),
                  'new_n': list(new_n.values()),
                  'title': list(new_n.keys())})
        lit.set_index(keys='n', inplace=True)
    print(lit)

    doc = docx.Document(inp)
    for p in paragraph_iterator(doc):  # doc.paragraphs:
        ss = t.search(p.text)
        while ss:
            l = ss[1]
            new_text = pack_ref(lit.loc[map(int, unpack_ref(l)[1])]['new_n'])
            print(ss[1], '->', new_text)
            replace_text_in_runs(p.runs, ss.span(1)[0],
                                 ss.span(1)[1], new_text)
            ss = t.search(p.text, pos=ss.span(1)[1]+1)
    doc.save(output)
    lit.drop_duplicates(subset='new_n', inplace=True)
    lit['n'] = lit['new_n']
    lit = lit.loc[lit.title != 'Not found']
    lit.to_excel(refs_output, columns=['n', 'title'], index=False)


if __name__ == '__main__':
    from argparse import ArgumentParser
    parser = ArgumentParser()
    parser.add_argument('input', help='Input docx file')
    parser.add_argument(
        '--refs', default=None, help='xlsx file with references. Must have cols: n and title')
    parser.add_argument('--start', '-n', default=1,
                        help='Start number to renumber references', type=int)
    parser.add_argument('--new_refs', default='new_lit.xlsx',
                        help='Output xls file')
    parser.add_argument('--output', '-o', default='rez.docx',
                        help='Output docx file')
    args = parser.parse_args()
    renumber_refs(args.input, args.output, args.refs,
                  args.new_refs, args.start)
