# %%
from colorama import Fore, Back, init
import docx
import re
from doctools_lib import replace_text_in_runs, paragraph_iterator

formula = re.compile(r'\(([\d.]+)\)')
stop_tag = re.compile(r'<e?stop>')
continue_tag = re.compile(r'<e?continue>')


def get_eq_params(input_str):
    p1 = re.compile(r'<e>(.*)<\\e>')
    start = re.compile(r'start\s*=\s*([\d]+)')
    prefix = re.compile(r'prefix\s*=\s*([\d.]+)')
    m1 = p1.search(input_str)
    if not m1:
        return None
    kws = {}
    kws['prefix'] = prefix.search(m1[1])[1] if prefix.search(m1[1]) else ''
    kws['start'] = int(start.search(m1[1])[1]) if start.search(m1[1]) else 1
    return kws


def analize_formulas(document, prefix='', start=1):
    pat = re.compile(r'[^а-яА-Я]*\((?P<num>[\d.-]+)\)\s*$')
    rez = {}
    do_analysis = True
    for p in paragraph_iterator(document):
        if stop_tag.search(p.text):
            do_analysis = False
        if continue_tag.search(p.text):
            do_analysis = True
        if not do_analysis:
            continue
        kw = get_eq_params(p.text)
        if kw:
            start = kw['start']
            prefix = kw['prefix']
            continue
        m = pat.match(p.text)
        if m:
            n = m.group('num')
            if not n in rez:
                rez[n] = [prefix, start]
                start += 1
    return rez


def renumber_formulas(inp, output, prefix, start):
    init(autoreset=True)
    formulas = analize_formulas(docx.Document(inp), prefix, start)
    doc = docx.Document(inp)
    do_replacement = True
    for p in paragraph_iterator(doc):  # doc.paragraphs:
        if stop_tag.search(p.text):
            do_replacement = False
            print(Fore.BLUE+'Stop equation renumbering...')
        if continue_tag.search(p.text):
            do_replacement = True
        if not do_replacement:
            continue
        ss = formula.search(p.text)
        while ss:
            n = ss[1]
            if not (n in formulas):
                print(
                    Fore.RED+f'Equation reference {n} is not found in document')
            else:
                new_str = f'{formulas[n][0]}{formulas[n][1]}'
                print(ss[1], '->', new_str)
                replace_text_in_runs(p.runs, ss.span(1)[0],
                                     ss.span(1)[1], new_str)
            ss = formula.search(p.text, pos=ss.span(1)[1]+1)
    doc.save(output)


if __name__ == '__main__':
    from argparse import ArgumentParser
    parser = ArgumentParser()
    parser.add_argument('input', help='Input docx file')
    parser.add_argument('--start', '-n', default=1,
                        help='Start number for figures numbering', type=int)
    parser.add_argument('--prefix', '-p', default='',
                        help='Default prefix for numbers')
    parser.add_argument(
        '--output', '-o', default='rez.docx', help='Output file')
    args = parser.parse_args()
    renumber_formulas(args.input, args.output, args.prefix, args.start)
